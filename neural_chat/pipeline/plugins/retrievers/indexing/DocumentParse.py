"""Wrapper for parsing the uploaded user file and then make document indexing."""

from langchain.vectorstores.chroma import Chroma
import pandas as pd
import PyPDF2
from docx import Document as DDocument
import re, json
import os
import unicodedata
from langchain.document_loaders import TextLoader, UnstructuredMarkdownLoader
from langchain.docstore.document import Document
from langchain.embeddings import HuggingFaceEmbeddings, HuggingFaceInstructEmbeddings
from langchain.vectorstores import Chroma
from haystack.schema import Document as SDocument


def uni_pro(text):
    """Check if the character is ASCII or falls in the category of non-spacing marks."""
    normalized_text = unicodedata.normalize('NFKD', text)
    filtered_text = ''
    for char in normalized_text:
        if ord(char) < 128 or unicodedata.category(char) == 'Mn':
            filtered_text += char
    return filtered_text


def load_pdf(pdf_path):
    """Read the pdf file."""
    pdf_file = open(pdf_path, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    text = ''
    for num in range(len(pdf_reader.pages)):
      page = pdf_reader.pages[num]
      text += page.extract_text()
    return text


def read_html(html_path):
    """Read the html file."""
    with open(html_path, 'r') as file:
        html = file.read()
    soup = BeautifulSoup(html, 'html.parser')
    text = soup.get_text(strip=True)
    return text


def read_txt(txt_path):
    """Read txt file."""
    with open(txt_path, 'r') as file:
        text = file.read()
    return text


def read_docx(doc_path):
    """Read docx file."""
    doc = DDocument(doc_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text
    return text


def read_md(md_path):
    """Read docx file."""
    loader = UnstructuredMarkdownLoader("instruction_data.md")
    text = loader.load()[0].page_content
    return text


def load_json(input, process, max_length):
    """Load and process json file."""
    data = []
    with open(input, 'r') as file:
        for line in file:
            json_obj = json.loads(line)
            data.append(json_obj)
    
    new_sens = []
    new_collec = []
    paragraphs = []
    for sub in data:
        sub['doc'].replace('#', " ")
        if not process:
            sub['doc'] = re.sub(r'\s+', ' ', sub['doc'])
            new_doc = [sub['doc'], sub['doc_id']]
            new_collect.append(new_doc)
        else:
            for sub in data:
                sub['doc'].replace('#', " ")
                split_sen = re.split(r'[.?!]', sub['doc'])
                for num in range(len(split_sen)):
                    split_sen[num] = re.sub(r'\s+', ' ', split_sen[num])
                    if num+1 < len(split_sen):
                        if len(split_sen[num])>max_length:
                            new_sens.append(split_sen[num].strip())
                        else:
                            split_sen[num+1]=split_sen[num]+split_sen[num+1]
                    else:
                        new_sens.append(split_sen[num])

            paragraphs = list(set(new_sens))
            for paragraph in paragraphs:
                new_doc = [paragraph, sub['doc_id']]
                new_collect.append(new_doc)
    return new_collect


def load_xlsx(input):
    df = pd.read_excel(input)
    all_data = []
    documents = []
    
    for index, row in df.iterrows():
        sub = "User Query: " + row['Questions'] + "Answer: " + row["Answers"]
        all_data.append(sub)

    for data in all_data:
        data.replace('#', " ")
        data = re.sub(r'\s+', ' ', data)
        new_doc = [data, input]
        documents.append(new_doc)
    return documents
    

def load_unstructured_data(input):
    """Load unstructured context."""
    if input.endswith("pdf"):
        text = load_pdf(input)
    elif input.endswith("docx"):
        text = read_docx(input)
    elif input.endswith("html"):
        text = read_html(input)
    elif input.endswith("txt"):
        text = read_txt(input)
    elif input.endswith("md"):
        text = read_md(input)
    
    text = text.replace('\n', '')
    text = text.replace('\n\n', '')
    text = uni_pro(text)
    text = re.sub(r'\s+', ' ', text)
    return text


def laod_structured_data(input, process, max_length):
    """Load structured context."""
    if input.endswith("jsonl"):
        content = load_json(input, process, max_length)
    else:
        content = load_xlsx(input)
    return content
        

def get_chuck_data(content, max_length, input):
    """Process the context to make it maintain a suitable length for the generation."""
    sentences = re.split('(?<=[;!.?])', content)

    paragraphs = []
    current_length = 0
    count = 0
    current_paragraph = ""
    for sub_sen in sentences:
        count +=1
        sentence_length = len(sub_sen)
        if current_length + sentence_length <= max_length:
            current_paragraph += sub_sen
            current_length += sentence_length
            if count == len(sentences):
                paragraphs.append([current_paragraph.strip(),input])
        else:
            paragraphs.append([current_paragraph.strip(),input])
            current_paragraph = sub_sen
            current_length = sentence_length

    return paragraphs


class DocumentIndexing:
    def __init__(self, retrieval_type="dense", document_store=None, output_dir="./output", process=False, embedding_model="hkunlp/instructor-large", max_length=378):
        """
        Wrapper for document indexing. Support dense and sparse indexing method.
        """
        self.retrieval_type = retrieval_type
        self.document_store = document_store
        self.process = process
        self.output_dir = output_dir
        self.embedding_model = embedding_model
        self.max_length = max_length
        
        
    def parse_document(self, input):
        """
        Parse the uploaded file.
        """
        if input.endswith("pdf") or input.endswith("docx") or input.endswith("html") or input.endswith("txt") or input.endswith("md"):
            content = load_unstructured_data(input)
            if self.process:
                chuck = get_chuck_data(content, self.max_length, input)
            else:
                chuck = [[content.strip(),input]]
        elif input.endswith("jsonl") or input.endswith("xlsx"):
            chuck = laod_structured_data(input, self.process, self.max_length)
        else:
            print("This file is ignored. Will support this file format soon.")
        return chuck
        
    
    def batch_parse_document(self, input):
        """
        Parse the uploaded batch files in the input folder.
        """
        paragraphs = []
        for dirpath, dirnames, filenames in os.walk(input):
            for filename in filenames:
                if filename.endswith("pdf") or filename.endswith("docx") or filename.endswith("html") or filename.endswith(
                        "txt") or filename.endswith("md"):
                    content = load_unstructured_data(os.path.join(dirpath, filename))
                    if self.process:
                        chuck = get_chuck_data(content, self.max_length, input)
                    else:
                        chuck = [[content.strip(),input]]
                    paragraphs += chuck
                elif filename.endswith("jsonl") or filename.endswith("xlsx"):
                    chuck = laod_structured_data(os.path.join(dirpath, filename), self.process, self.max_length)
                    paragraphs += chuck
                else:
                    print("This file {} is ignored. Will support this file format soon.".format(filename))
        return paragraphs
    
    
    def KB_construct(self, input):
        """
        Construct the local knowledge base based on the uploaded file/files.
        """
        if self.retrieval_type == "dense":
            if os.path.exists(input):
                if os.path.isfile(input):
                    data_collection = self.parse_document(input)
                elif os.path.isdir(input):
                    data_collection = self.batch_parse_document(input)
                else:
                    print("Please check your upload file and try again!")
                    
                documents = []
                for data, meta in data_collection:
                    metadata = {"source": meta}
                    new_doc = Document(page_content=data, metadata=metadata)
                    documents.append(new_doc)
                embedding = HuggingFaceInstructEmbeddings(model_name=self.embedding_model)
                vectordb = Chroma.from_documents(documents=documents, embedding=embedding,
                                                 persist_directory=self.output_dir)
                vectordb.persist()
                print("success")
                return vectordb
            else:
                print("There might be some errors, please wait and try again!")
        else:
            if os.path.exists(input):
                if os.path.isfile(input):
                    data_collection = self.parse_document(input)
                elif os.path.isdir(input):
                    data_collection = self.batch_parse_document(input)
                else:
                    print("Please check your upload file and try again!")
                if self.document_store == "inmemory":
                    document_store = InMemoryDocumentStore(use_gpu=False, use_bm25=True)
                elif self.document_stor == "Elasticsearch":
                    document_store = ElasticsearchDocumentStore(host="localhost", index="elastic_index_1",
                                                                port=9200, search_fields=["content", "title"])

                documents = []
                for data, meta in data_collection:
                    metadata = {"source": meta}
                    new_doc = SDocument(content=data, metadata=metadata)
                    documents.append(new_doc)
                document_store.write_documents(documents)
                print("success")
                return document_store
            else:
                print("There might be some errors, please wait and try again!")