import os
import argparse
from langchain.embeddings import HuggingFaceEmbeddings, HuggingFaceInstructEmbeddings
from langchain.vectorstores import Chroma
from chromadb.utils import embedding_functions
from langchain.docstore.document import Document
import re, json
from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownTextSplitter
from langchain.document_loaders import TextLoader, UnstructuredMarkdownLoader
import PyPDF2
from haystack.schema import Document as SDocument
from docx import Document as DDocument
from haystack.document_stores import ElasticsearchDocumentStore, InMemoryDocumentStore
import pandas as pd


def split_paragraph(text, jsonl_name, max_length=378):
    new_sens = []
    documents = []
    for sub in text:
        sub['doc'].replace('#', " ")
        sub['doc'] = re.sub(r'\s+', ' ', sub['doc'])
        new_doc = Document(page_content=sub['doc'], metadata={"source": sub['doc_id']})
        documents.append(new_doc)
    return documents


## indexing for jsonl file
def d_load_jsonl_file(file_path, process, max_length=378):
    data = []
    with open(file_path, 'r') as file:
        for line in file:
            json_obj = json.loads(line)
            data.append(json_obj)

    new_sens = []
    documents = []
    paragraphs = []
    for sub in data:
        sub['doc'].replace('#', " ")
        if not process:
            sub['doc'] = re.sub(r'\s+', ' ', sub['doc'])
            new_doc = Document(page_content=sub['doc'], metadata={"source": sub['doc_id']})
            documents.append(new_doc)
        else:
            for sub in data:
                sub['doc'].replace('#', " ")
                split_sen = re.split(r'[.?!]', sub['doc'])
                for num in range(len(split_sen)):
                    split_sen[num] = re.sub(r'\s+', ' ', split_sen[num])
                    if num+1 < len(split_sen):
                        if len(split_sen[num])>max_length:
                            new_sens.append(split_sen[num].strip())
                        else:
                            split_sen[num+1]=split_sen[num]+split_sen[num+1]
                    else:
                        new_sens.append(split_sen[num])

            print("length for origin", len(new_sens))
            paragraphs = list(set(new_sens))
            print("length for processed", len(new_sens))
            documents = []
            metadata = {"source": file_path}
            for paragraph in paragraphs:
                new_doc = Document(page_content=paragraph, metadata=metadata)
                documents.append(new_doc)
    return documents


# def d_load_xlsx_file(file_path, process, max_length=378):
#     data = []
#     data = pd.read_excel(file_path)
# 
#     new_sens = []
#     documents = []
#     paragraphs = []
#     for sub in data:
#         sub['doc'].replace('#', " ")
#         if not process:
#             sub['doc'] = re.sub(r'\s+', ' ', sub['doc'])
#             new_doc = Document(page_content=sub['doc'], metadata={"source": sub['doc_id']})
#             documents.append(new_doc)

## indexing for pdf file
def d_load_file(file_path, process, max_length=378):
    if file_path.endswith("pdf"):
        text = load_pdf(file_path)
    elif file_path.endswith("docx"):
        text = read_docx(file_path)

    text = text.replace('\n', '')
    text = text.replace('\n\n', '')
    text = re.sub(r'\s+', ' ', text)
    """
    split the document
    """
    sentences = re.split('(?<=[;!.?])', text)

    new_sents = []
    for i in range(int(len(sentences) / 2)):
        sent = sentences[2 * i] + sentences[2 * i + 1]
        new_sents.append(sent)
    if len(sentences) % 2 == 1:
        new_sents.append(sentences[len(sentences) - 1])

    paragraphs = []
    current_length = 0
    current_paragraph = ""
    for sentence in new_sents:
        sentence_length = len(sentence)
        if current_length + sentence_length <= max_length:
            current_paragraph += sentence
            current_length += sentence_length
        else:
            paragraphs.append(current_paragraph.strip())
            current_paragraph = sentence
            current_length = sentence_length
    print("length for origin", len(paragraphs))
    paragraphs.append(current_paragraph.strip())
    paragraphs = list(set(paragraphs))
    print("length for processed", len(paragraphs))
    documents = []
    metadata = {"source": file_path}
    for paragraph in paragraphs:
        new_doc = Document(page_content=paragraph, metadata=metadata)
        documents.append(new_doc)
    return documents


### Load with spare embedding for jsonl file
def s_load_jsonl_file(file_path, process, document_store, max_length=378):
    data = []
    with open(file_path, 'r') as file:
        for line in file:
            json_obj = json.loads(line)
            data.append(json_obj)

    new_sens = []
    documents = []
    paragraphs = []
    for sub in data:
        sub['doc'].replace('#', " ")
        if not process:
            sub['doc'] = re.sub(r'\s+', ' ', sub['doc'])
            new_doc = SDocument(content=sub['doc'], meta={"source": sub['doc_id']})
            documents.append(new_doc)
        else:
            for sub in data:
                sub['doc'].replace('#', " ")
                split_sen = re.split(r'[.?!]', sub['doc'])
                for num in range(len(split_sen)):
                    split_sen[num] = re.sub(r'\s+', ' ', split_sen[num])
                    if num+1 < len(split_sen):
                        if len(split_sen[num])>max_length:
                            new_sens.append(split_sen[num].strip())
                        else:
                            split_sen[num+1]=split_sen[num]+split_sen[num+1]
                    else:
                        new_sens.append(split_sen[num])

            print("length for origin", len(new_sens))
            paragraphs = list(set(new_sens))
            print("length for processed", len(new_sens))
            documents = []
            metadata = {"source": file_path}
            for paragraph in paragraphs:
                new_doc = SDocument(content=paragraph, meta=metadata)
                documents.append(new_doc)
    document_store.write_documents(documents)
    return document_store


### Load with spare embedding for pdf file
def s_load_file(file_path, process, document_store, max_length=378):
    if file_path.endswith("pdf"):
        text = load_pdf(file_path)
    elif file_path.endswith("docx"):
        text = read_docx(file_path)

    text = text.replace('\n', '')
    text = text.replace('\n\n', '')
    text = re.sub(r'\s+', ' ', text)
    """
    split the document
    """
    sentences = re.split('(?<=[;!.?])', text)

    new_sents = []
    for i in range(int(len(sentences) / 2)):
        sent = sentences[2 * i] + sentences[2 * i + 1]
        new_sents.append(sent.strip())
    if len(sentences) % 2 == 1:
        new_sents.append(sentences[len(sentences) - 1])

    paragraphs = []
    current_length = 0
    current_paragraph = ""
    for sentence in new_sents:
        sentence_length = len(sentence)
        if current_length + sentence_length <= max_length:
            current_paragraph += sentence
            current_length += sentence_length
        else:
            paragraphs.append(current_paragraph.strip())
            current_paragraph = sentence
            current_length = sentence_length
    print("length for origin", len(paragraphs))
    paragraphs.append(current_paragraph.strip())
    paragraphs = list(set(paragraphs))
    print("length for processed", len(paragraphs))
    documents = []
    metadata = {"source": file_path}
    for paragraph in paragraphs:
        new_doc = SDocument(content=paragraph, metadata=metadata)
        documents.append(new_doc)
    document_store.write_documents(documents)
    
    return document_store


def persist_embedding(documents, persist_directory, model_path):
    ## persistly save the local file into disc
    embedding = HuggingFaceInstructEmbeddings(model_name=model_path)
    vectordb = Chroma.from_documents(documents=documents, embedding=embedding, persist_directory=persist_directory)
    vectordb.persist()
    vectordb = None

    
def read_docx(doc_path):
    doc = DDocument(doc_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text
    return text

def load_pdf(pdf_path):
    pdf_file = open(pdf_path, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    text = ''
    for num in range(len(pdf_reader.pages)):
      page = pdf_reader.pages[num]
      text += page.extract_text()
    return text


if __name__ == "__main__":

    parser = argparse.ArgumentParser()
    parser.add_argument('--file_path', type=str, help='The user upload file.',
                        default="/data1/lkk/llm_inference/chat-langchain/inc_documents_formated.jsonl")
    parser.add_argument('--process', type=bool,
                        help='Whether or not to proceed the load content.',
                        default=False)
    parser.add_argument('--embedding_model', type=str, help='Select which model to embed the content.', default='/data1/lkk/instructor_large/')
    parser.add_argument('--output_path', type=str, help='Where to save the embedding.', default='db_jsonl122')
    parser.add_argument('--embedding_method', type=str, help='Select to use dense retrieval or sparse retrieval.', default='dense')
    parser.add_argument('--store', type=str, help='Select to use dense retrieval or sparse retrieval.',
                        default='dense')

    args = parser.parse_args()

    if args.embedding_method == "dense":  # currently use Chroma as the dense retrieval datastore
        if args.file_path.endswith("jsonl"):
            documents = d_load_jsonl_file(args.file_path, args.process)
        elif args.file_path.endswith("pdf") or args.file_path.endswith("docx"):
            documents = d_load_file(args.file_path, args.process)
        else:
            print("{} is ignored. Will support this file format soon.".format(args.file_path))
        persist_embedding(documents, args.output_path, args.embedding_model)
    elif args.embedding_method == "sparse":   # sparse retrieval datastores has inmemory and Elasticsearch
        if args.store == "inmemory":
            document_store = InMemoryDocumentStore(use_gpu=False, use_bm25=True)
        elif args.store == "Elasticsearch":
            document_store = ElasticsearchDocumentStore(host="localhost", index="elastic_index_1",
                                               port=9200, search_fields=["content", "title"])
        # import pdb;pdb.set_trace()
        if args.file_path.endswith("jsonl"):
            document_store = s_load_jsonl_file(args.file_path, args.process, document_store)
        elif args.file_path.endswith("pdf") or args.file_path.endswith("docx"):
            document_store = s_load_file(args.file_path, args.process, document_store)
        else:
            print("{} is ignored. Will support this file format soon.".format(args.file_path))
                
